# frontend/app.py
import io
import streamlit as st
from docx import Document
from api_client import check_health, generate_content, get_tones

st.set_page_config(
    page_title="Content Generator",
    page_icon="✍️",
    layout="wide",
)


def build_docx(result: dict) -> bytes:
    """Build a Word document from the generated content."""
    doc = Document()

    doc.add_heading("Content Package", 0)
    doc.add_paragraph(f"Topic: {result['topic']}")
    doc.add_paragraph(f"Tone: {result['tone']}")
    doc.add_paragraph(f"Audience: {result['audience']}")

    doc.add_heading("Blog Post Outline", level=1)
    doc.add_paragraph(result["outline"])

    doc.add_heading("Blog Post", level=1)
    doc.add_paragraph(result["blog_post"])

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(result["summary"])

    doc.add_heading("Social Media Captions", level=1)
    doc.add_heading("LinkedIn", level=2)
    doc.add_paragraph(result["captions"]["linkedin"])
    doc.add_heading("Twitter / X", level=2)
    doc.add_paragraph(result["captions"]["twitter"])
    doc.add_heading("Instagram", level=2)
    doc.add_paragraph(result["captions"]["instagram"])

    doc.add_heading("Email Copy", level=1)
    doc.add_paragraph(f"Subject: {result['email']['subject']}")
    doc.add_paragraph(f"Preview: {result['email']['preview']}")
    doc.add_heading("Body", level=2)
    doc.add_paragraph(result["email"]["body"])
    doc.add_paragraph(f"CTA Button: {result['email']['cta']}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("✍️ Content Generator")
    st.caption("AI-powered content pipeline")

    st.divider()

    is_healthy = check_health()
    if is_healthy:
        st.success("API connected", icon="🟢")
    else:
        st.error("API unreachable", icon="🔴")

    st.divider()
    st.subheader("Settings")

    # Load tones from API
    tones_result = get_tones()
    if tones_result["success"]:
        tone_options = {
            t["value"].title(): t["value"]
            for t in tones_result["data"]
        }
    else:
        tone_options = {"Professional": "professional", "Casual": "casual"}

    selected_tone_label = st.selectbox(
        "Writing Tone",
        options=list(tone_options.keys()),
        help="Sets the style and voice of all generated content",
    )
    selected_tone = tone_options[selected_tone_label]

    audience = st.text_input(
        "Target Audience",
        value="general audience",
        help="Who is this content written for?",
        placeholder="e.g. Nigerian entrepreneurs",
    )

    st.divider()
    st.caption("⏱️ Generation takes ~20 seconds")
    st.caption("Five LLM calls run sequentially")


# ── Main area ─────────────────────────────────────────────────────────────────
st.title("Content Generation Pipeline")
st.caption(
    "Enter a topic and get a complete content package: "
    "blog post, social media captions, and email copy."
)

if not is_healthy:
    st.warning("⚠️ Backend API is not reachable. Start your FastAPI server first.")
    st.stop()

st.divider()

# ── Topic input ───────────────────────────────────────────────────────────────
topic = st.text_area(
    "What do you want to write about?",
    height=100,
    placeholder="e.g. How AI is helping small businesses in Nigeria",
    help="Be specific — better topics produce better content",
)

col1, col2 = st.columns([1, 4])
with col1:
    generate_clicked = st.button(
        "Generate Content",
        type="primary",
        disabled=not topic.strip(),
        use_container_width=True,
    )

if generate_clicked and topic.strip():
    with st.spinner("Running 5-step content pipeline..."):
        result = generate_content(
            topic=topic,
            tone=selected_tone,
            audience=audience,
        )

    if result["success"]:
        data = result["data"]
        st.success("✅ Content generated successfully")

        # Download button
        docx_bytes = build_docx(data)
        st.download_button(
            label="⬇️ Download as Word Document",
            data=docx_bytes,
            file_name=f"content_{topic[:30].replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        st.divider()

        # Content tabs
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "📋 Outline",
            "📝 Blog Post",
            "📱 Social Captions",
            "📧 Email Copy",
            "💡 Summary",
        ])

        with tab1:
            st.markdown("### Blog Post Outline")
            st.markdown(data["outline"])

        with tab2:
            st.markdown("### Blog Post")
            st.markdown(data["blog_post"])
            st.caption(f"{len(data['blog_post'].split())} words")

        with tab3:
            st.markdown("### Social Media Captions")

            st.markdown("**LinkedIn**")
            st.info(data["captions"]["linkedin"])

            st.markdown("**Twitter / X**")
            st.info(data["captions"]["twitter"])

            st.markdown("**Instagram**")
            st.info(data["captions"]["instagram"])

        with tab4:
            st.markdown("### Email Copy")
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Subject Line", data["email"]["subject"])
            with col_b:
                st.metric("CTA Button", data["email"]["cta"])
            st.markdown("**Preview Text:**")
            st.caption(data["email"]["preview"])
            st.markdown("**Email Body:**")
            st.markdown(data["email"]["body"])

        with tab5:
            st.markdown("### Content Summary")
            st.info(data["summary"])

    else:
        st.error(f"❌ {result['error']}")